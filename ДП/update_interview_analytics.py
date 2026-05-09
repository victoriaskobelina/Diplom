from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


BASE = Path(r"C:\Users\1\Desktop\D_1\ДП")
DOCX = BASE / "Интервью.docx"
COPY = BASE / "Интервью_нумерация_исправлена.docx"

OLD_QUESTION = "Какие статистические данные и аналитика должны быть доступны в системе?"
NEW_QUESTION = "Какие статистические данные и аналитика должны быть доступны для каждой роли?"
NEW_ANSWER = (
    "Для студента должна быть доступна аналитика по собственным результатам: список пройденных тестов, "
    "количество использованных попыток, набранные баллы, процент выполнения, оценка и дата прохождения. "
    "Для преподавателя должна быть доступна аналитика по его тестам и учебным группам: список студентов, "
    "список назначенных тестов, наличие или отсутствие результата по каждому студенту, средний балл по тесту "
    "и группе, процент выполнения, количество прошедших и не прошедших тестирование, а также выгрузка этих "
    "данных в Excel. Для администратора должна быть доступна общая статистика по системе: количество студентов, "
    "преподавателей, учебных групп, дисциплин, тестов, завершенных попыток, а также сведения из журнала действий "
    "для контроля активности пользователей."
)


def normalize(value: str) -> str:
    return " ".join(value.split()).strip()


def format_cell(cell, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(12)


def set_cell_text(cell, text: str, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = text
    format_cell(cell, align)


def main():
    doc = Document(DOCX)
    table = doc.tables[0]

    updated = False
    for row in table.rows[1:]:
        question = normalize(row.cells[1].text)
        if question in {normalize(OLD_QUESTION), normalize(NEW_QUESTION)}:
            set_cell_text(row.cells[1], NEW_QUESTION)
            set_cell_text(row.cells[2], NEW_ANSWER)
            updated = True
            break

    if not updated:
        raise RuntimeError("Не найден вопрос про аналитику и статистические данные.")

    for index, row in enumerate(table.rows[1:], start=1):
        set_cell_text(row.cells[0], str(index), WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(DOCX)
    shutil.copy2(DOCX, COPY)

    nums = [row.cells[0].text.strip() for row in Document(DOCX).tables[0].rows[1:]]
    expected = [str(i) for i in range(1, len(nums) + 1)]
    print(DOCX)
    print("question_count=", len(nums))
    print("numbers_ok=", nums == expected)


if __name__ == "__main__":
    main()
