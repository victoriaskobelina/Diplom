from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


BASE = Path(r"C:\Users\1\Desktop\D_1\ДП")
DOCX = BASE / "Интервью.docx"
BACKUP = BASE / "Интервью_до_доработки.docx"


NEW_ROWS = [
    {
        "after": "Должна ли быть возможность добавления изображения к вопросу?",
        "question": "Как должны начисляться баллы за ответы в тесте?",
        "answer": (
            "Один правильный ответ на вопрос должен равняться одному баллу. "
            "Все вопросы имеют одинаковый вес, отдельная градация баллов по вопросам не требуется. "
            "Итоговый результат рассчитывается как количество правильных ответов, а также может отображаться в процентах и оценке."
        ),
    },
    {
        "after": "Нужна ли выгрузка результатов тестирования?",
        "question": "Нужна ли выгрузка результатов тестирования по учебным группам в Excel?",
        "answer": (
            "Да. Преподавателю необходима выгрузка результатов по выбранной учебной группе в формате Excel. "
            "В файле должен отображаться список студентов группы, список назначенных тестов и заполняемость результатами: "
            "для каждого студента должно быть видно, по каким тестам результат уже есть, а по каким тестирование еще не пройдено."
        ),
    },
    {
        "after": "Нужна ли выгрузка результатов тестирования по учебным группам в Excel?",
        "question": "Какие статистические данные и аналитика должны быть доступны в системе?",
        "answer": (
            "Необходимо отображать количество студентов, прошедших и не прошедших тестирование, средний результат по тесту и группе, "
            "процент выполнения, оценки, количество использованных попыток, а также список студентов без результата. "
            "Эти данные должны помогать преподавателю анализировать успеваемость группы и выявлять проблемные темы."
        ),
    },
    {
        "after": "Как будет проверяться готовность программного продукта?",
        "question": "Какая документация должна быть подготовлена по итогам разработки?",
        "answer": (
            "По итогам разработки необходимо подготовить техническое задание, руководство пользователя для основных ролей "
            "системы, руководство администратора, а также краткую инструкцию по установке и эксплуатации веб-системы на локальном сервере."
        ),
    },
]


def normalize_text(value: str) -> str:
    return " ".join(value.split()).strip()


def set_cell_text(cell, text: str, *, align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = text
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(12)


def find_row_index(table, question: str) -> int:
    target = normalize_text(question)
    for index, row in enumerate(table.rows):
        if len(row.cells) > 1 and normalize_text(row.cells[1].text) == target:
            return index
    raise ValueError(f"Не найден вопрос: {question}")


def insert_after(table, after_question: str, question: str, answer: str):
    index = find_row_index(table, after_question)
    new_tr = deepcopy(table.rows[index]._tr)
    table.rows[index]._tr.addnext(new_tr)
    new_row = table.rows[index + 1]
    set_cell_text(new_row.cells[0], "", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(new_row.cells[1], question)
    set_cell_text(new_row.cells[2], answer)


def renumber_questions(table):
    for number, row in enumerate(table.rows[1:], start=1):
        set_cell_text(row.cells[0], str(number), align=WD_ALIGN_PARAGRAPH.CENTER)


def main():
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)

    doc = Document(DOCX)
    table = doc.tables[0]

    existing_questions = {
        normalize_text(row.cells[1].text)
        for row in table.rows[1:]
        if len(row.cells) > 1
    }

    for item in NEW_ROWS:
        if normalize_text(item["question"]) not in existing_questions:
            insert_after(table, item["after"], item["question"], item["answer"])
            existing_questions.add(normalize_text(item["question"]))

    renumber_questions(table)
    doc.save(DOCX)
    print(DOCX)
    print(BACKUP)


if __name__ == "__main__":
    main()
