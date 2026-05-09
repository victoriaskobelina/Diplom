from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


BASE = Path(r"C:\Users\1\Desktop\D_1\ДП")
DOCX = BASE / "А.docx"
BACKUP = BASE / "А_до_разделения_31_вопроса.docx"

TARGET_QUESTION = "Какие статистические данные и аналитика должны быть доступны для каждой роли?"

SPLIT_ROWS = [
    (
        "Какие статистические данные и аналитика должны быть доступны студенту?",
        "Студенту должна быть доступна аналитика по собственным результатам: список пройденных тестов, "
        "количество использованных попыток, набранные баллы, процент выполнения, оценка и дата прохождения. "
        "Также студент должен видеть, какие тесты уже пройдены, а какие еще доступны для прохождения.",
    ),
    (
        "Какие статистические данные и аналитика должны быть доступны преподавателю?",
        "Преподавателю должна быть доступна аналитика по его тестам и учебным группам: список студентов, "
        "список назначенных тестов, наличие или отсутствие результата по каждому студенту, средний балл по тесту "
        "и группе, процент выполнения, количество прошедших и не прошедших тестирование, количество использованных "
        "попыток, а также выгрузка этих данных в Excel.",
    ),
    (
        "Какие статистические данные и аналитика должны быть доступны администратору?",
        "Администратору должна быть доступна общая статистика по системе: количество студентов, преподавателей, "
        "учебных групп, дисциплин, тестов и завершенных попыток. Также администратор должен видеть сведения из "
        "журнала действий для контроля активности пользователей и администрирования платформы.",
    ),
]


def normalize(value: str) -> str:
    return " ".join(value.split()).strip()


def format_cell(cell, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = align
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.left_indent = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.15
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(12)


def set_cell(cell, text: str, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = text
    format_cell(cell, align)


def find_row(table, question: str) -> int:
    target = normalize(question)
    for index, row in enumerate(table.rows):
        if len(row.cells) > 1 and normalize(row.cells[1].text) == target:
            return index
    raise RuntimeError(f"Не найден вопрос: {question}")


def main():
    if not BACKUP.exists():
        shutil.copy2(DOCX, BACKUP)

    doc = Document(DOCX)
    table = doc.tables[0]

    # If the question was already split, do not duplicate rows; just renumber.
    existing_questions = {normalize(row.cells[1].text) for row in table.rows[1:]}
    if normalize(SPLIT_ROWS[1][0]) not in existing_questions:
        row_index = find_row(table, TARGET_QUESTION)
        set_cell(table.rows[row_index].cells[1], SPLIT_ROWS[0][0])
        set_cell(table.rows[row_index].cells[2], SPLIT_ROWS[0][1])

        insert_after = row_index
        for question, answer in SPLIT_ROWS[1:]:
            new_tr = deepcopy(table.rows[insert_after]._tr)
            table.rows[insert_after]._tr.addnext(new_tr)
            insert_after += 1
            new_row = table.rows[insert_after]
            set_cell(new_row.cells[0], "", WD_ALIGN_PARAGRAPH.CENTER)
            set_cell(new_row.cells[1], question)
            set_cell(new_row.cells[2], answer)

    for number, row in enumerate(table.rows[1:], start=1):
        set_cell(row.cells[0], str(number), WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(DOCX)

    check = Document(DOCX)
    nums = [row.cells[0].text.strip() for row in check.tables[0].rows[1:]]
    expected = [str(i) for i in range(1, len(nums) + 1)]
    print(DOCX)
    print("question_count=", len(nums))
    print("numbers_ok=", nums == expected)
    print(BACKUP)


if __name__ == "__main__":
    main()
